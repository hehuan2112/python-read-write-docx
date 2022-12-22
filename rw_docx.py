'''
Read/write docx file example

In this script, we plan to generate several files based on an XLSX file.
We will use each row in the XLSX file to create one docx file.
'''
import re
import docx
import pandas as pd

def replace_variables_in_text(text, row):
    '''
    Replace the pattern {{variable_name}} with data in `row`
    '''
    ###################################################
    # The easy way of replacement
    # Just use the `in` operator to find a specific pattern 
    # For example, the {{name}} in the text.
    # If there are a few variables need to be replace
    # This way is easy to implement and understand.
    # In fact, you don't need
    ###################################################
    if '{{name}}' in text:
        if 'name' in row:
            print('* wow! found {{name}} and replace it with %s!' % (row['name']))
            # replace the big name
            text = text.replace('{{name}}', row['name'])
    
    ###################################################
    # The better way of replacement.
    # As we already know the pattern {{variable_name}},
    # We can use regex to search all possible variables
    # in the given text.
    # There are some template engines that can provide 
    # more powerful features, such jinja, mako, etc.
    ###################################################
    # This pattern may catch {{name}}, {{other_name}}, 
    # {{ with_space }}, {{ name_123 }}, etc.
    # As there may be space between the variable name and double curly bracket
    # the `\s*` is needed.
    # For more about the regex, please check https://regex101.com/
    # and other tutorials about the regex.
    regex = r"{{\s*(\S+)\s*}}"

    # find all patterns in text
    matches = re.finditer(regex, text, re.MULTILINE)

    # if any matches are found, replace thme
    for matchNum, match in enumerate(matches, start=1):
        # find the matched string
        # e.g., {{name}}, {{ address }}
        matched_str = match.group()

        # we know that there is one group in the regex match
        # so we can get the exact variable name
        # e.g., name, address
        variable_name = match.group(1)

        # then, let's do the replacement if this variable 
        if variable_name in row:
            text = text.replace(
                matched_str, 
                row[variable_name]
            )
            print('* regex found %s and replace it with %s!' % (
                matched_str, row[variable_name]
            ))
        else:
            print('* unknown pattern %s' % matched_str)

    return text

        
# read xlsx file for data
df = pd.read_excel('data.xlsx')
print('* loaded %s rows of data' % (len(df)))

# for each row, we want to create a seperate docx file
for idx, dfrow in df.iterrows():
    print('*'*60)
    # open the template
    doc = docx.Document("template.docx")
    print('* loaded template docx file')

    # search all paragraphs
    print('* found %s paragraphs ' % (len(doc.paragraphs)))
    for i in range(len(doc.paragraphs)):
        # for each paragraph, 
        doc.paragraphs[i].text = replace_variables_in_text(doc.paragraphs[i].text, dfrow)
        
    # search all tables
    print('* found %s tables ' % (len(doc.tables)))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    # the same process as in paragraphs
                    # it's better to define a function to process
                    p.text = replace_variables_in_text(p.text, dfrow)

    # save the result
    fn_output = f"output-{idx}.docx"
    doc.save(fn_output)
    print("* saved output file %s" % fn_output)

print('* done!')